#!/usr/bin/env python3
"""
attachment_extract - Extract text content from PDF, DOCX, XLSX attachments.

Usage:
    python attachment_extract.py --input INPUT --output OUTPUT [--format FORMAT]

Options:
    --input INPUT        Input file path (PDF, DOCX, XLSX)
    --output OUTPUT      Output JSON file for extracted content
    --format FORMAT      File format [auto-detect, pdf, docx, xlsx]

Example:
    python attachment_extract.py --input document.pdf --output content.json
"""

import argparse
import json
import sys
from pathlib import Path

try:
    import fitz  # PyMuPDF for PDF
except ImportError:
    fitz = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import openpyxl
except ImportError:
    openpyxl = None


def extract_pdf(file_path):
    """Extract text from PDF file."""
    if fitz is None:
        return {"error": "PyMuPDF not installed. Run: pip install pymupdf"}

    try:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return {
            "text": text,
            "pages": len(doc) if 'doc' in dir() else 0,
            "chars": len(text)
        }
    except Exception as e:
        return {"error": f"Failed to extract PDF: {str(e)}"}


def extract_docx(file_path):
    """Extract text from DOCX file."""
    if DocxDocument is None:
        return {"error": "python-docx not installed. Run: pip install python-docx"}

    try:
        doc = DocxDocument(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Extract from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    tables_text.append(" | ".join(row_text))

        full_text = "\n".join(paragraphs)
        if tables_text:
            full_text += "\n\n[Tables]\n" + "\n".join(tables_text)

        return {
            "text": full_text,
            "paragraphs": len(paragraphs),
            "tables": len(doc.tables),
            "chars": len(full_text)
        }
    except Exception as e:
        return {"error": f"Failed to extract DOCX: {str(e)}"}


def extract_xlsx(file_path):
    """Extract text from XLSX file."""
    if openpyxl is None:
        return {"error": "openpyxl not installed. Run: pip install openpyxl"}

    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)
        sheets_text = []

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            rows_text = []

            for row in sheet.iter_rows(values_only=True):
                row_values = [str(cell) if cell is not None else "" for cell in row]
                if any(v.strip() for v in row_values):
                    rows_text.append(" | ".join(row_values))

            if rows_text:
                sheets_text.append(f"[Sheet: {sheet_name}]\n" + "\n".join(rows_text))

        full_text = "\n\n".join(sheets_text)

        return {
            "text": full_text,
            "sheets": len(wb.sheetnames),
            "sheet_names": wb.sheetnames,
            "chars": len(full_text)
        }
    except Exception as e:
        return {"error": f"Failed to extract XLSX: {str(e)}"}


def detect_format(file_path):
    """Auto-detect file format from extension."""
    suffix = Path(file_path).suffix.lower()
    format_map = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".xlsx": "xlsx",
        ".xls": "xlsx",
    }
    return format_map.get(suffix, "unknown")


def extract_content(file_path, output_path, format_type="auto"):
    """Extract content from file and save to output."""
    path = Path(file_path)

    if not path.exists():
        return {"error": f"File not found: {file_path}"}

    if format_type == "auto":
        format_type = detect_format(file_path)

    if format_type == "unknown":
        return {"error": f"Unknown file format: {path.suffix}"}

    # Extract based on format
    if format_type == "pdf":
        result = extract_pdf(file_path)
    elif format_type == "docx":
        result = extract_docx(file_path)
    elif format_type == "xlsx":
        result = extract_xlsx(file_path)
    else:
        result = {"error": f"Unsupported format: {format_type}"}

    # Add metadata
    result["file_name"] = path.name
    result["file_type"] = format_type
    result["file_size"] = path.stat().st_size

    # Write output
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    return result


def main():
    parser = argparse.ArgumentParser(description="Extract content from PDF/DOCX/XLSX files")
    parser.add_argument("--input", required=True, help="Input file path")
    parser.add_argument("--output", required=True, help="Output JSON file")
    parser.add_argument("--format", default="auto",
                        choices=["auto", "pdf", "docx", "xlsx"],
                        help="File format [default: auto-detect]")

    args = parser.parse_args()

    result = extract_content(args.input, args.output, args.format)

    if "error" in result:
        print(f"Error: {result['error']}")
        sys.exit(1)
    else:
        print(f"Extracted {result.get('chars', 0)} chars from {result.get('file_name', 'file')}")
        print(f"Output saved to {args.output}")


if __name__ == "__main__":
    main()
